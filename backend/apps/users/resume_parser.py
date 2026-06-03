import json
import re
from io import BytesIO

from apps.tailor.services import (
    _extract_json_object,
    _generate_with_anthropic,
    _generate_with_openai,
    extract_pdf_text,
)

from .services import normalize_profile_sections

DEFAULT_PARSE_PROVIDER = "openai"
DEFAULT_PARSE_MODEL = "gpt-5.4-mini"
MAX_RESUME_BYTES = 10 * 1024 * 1024
PDF_CONTENT_TYPES = {"application/pdf"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
PDF_EXTENSIONS = (".pdf",)
DOCX_EXTENSIONS = (".docx",)
PROFILE_REVIEW_SECTIONS = ("experience", "projects", "education", "certificates", "skills")
DUPLICATE_REVIEW_ACTIONS = {"keep_existing", "use_parsed", "merge"}
NEW_REVIEW_ACTIONS = {"add", "skip"}

PROFILE_PARSER_SYSTEM_PROMPT = (
    "You are a resume parser. Extract truthful candidate profile data from resume text. "
    "Return only strict JSON with no markdown, comments, or code fences."
)

PROFILE_PARSER_USER_PROMPT = """
Extract this resume into the exact JSON schema below.

Rules:
- Return strict JSON only.
- Use empty strings or empty arrays for missing values.
- Do not infer facts not present in the resume.
- Keep dates as they appear, such as "Jan 2024" or "2020 - 2024".
- Put resume bullets into description arrays as plain text.
- For skills, return objects with name and category. Category may be empty.
- For projects, keep stack as an array of technologies.
- Ignore unsupported fields.

JSON schema:
{
  "full_name": "",
  "display_name": "",
  "basic": {
    "phone_country_code": "+1",
    "phone": "",
    "contact_email": "",
    "location": "",
    "headline": "",
    "github_url": "",
    "linkedin_url": "",
    "portfolio_url": "",
    "summary": ""
  },
  "experience": [
    {
      "company": "",
      "role": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "description": []
    }
  ],
  "projects": [
    {
      "name": "",
      "label": "",
      "stack": [],
      "description": [],
      "live_url": "",
      "github_url": "",
      "start_date": "",
      "end_date": ""
    }
  ],
  "education": [
    {
      "school": "",
      "degree": "",
      "major": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "gpa": "",
      "awards": [],
      "relevant_coursework": []
    }
  ],
  "certificates": [
    {
      "name": "",
      "issuer": "",
      "issue_date": "",
      "expiration_date": "",
      "credential_id": "",
      "credential_url": ""
    }
  ],
  "skills": [
    {
      "name": "",
      "category": ""
    }
  ]
}

Resume text:
%%RESUME_TEXT%%
""".strip()

PROFILE_REVIEW_SYSTEM_PROMPT = (
    "You compare parsed resume profile entries with an existing candidate profile. "
    "Return only strict JSON with no markdown, comments, or code fences."
)

PROFILE_REVIEW_USER_PROMPT = """
Compare every parsed profile entry against the existing profile entries.

Rules:
- Return strict JSON only.
- Review all parsed entries in these sections: experience, projects, education, certificates, skills.
- For each parsed entry, return exactly one review item.
- Use status "duplicate" when the parsed entry appears to describe the same real-world item as an existing entry, even if titles, names, dates, stacks, or bullet points differ.
- Use status "new" only when no existing entry describes the same real-world item.
- For duplicate items, include existing_index and recommend one of: keep_existing, use_parsed, merge.
- For new items, set existing_index to null and recommend one of: add, skip.
- Prefer "merge" for duplicates when both entries contain useful non-conflicting details or different bullets.
- Prefer "use_parsed" when the parsed resume entry appears more complete and current.
- Prefer "add" for new entries unless the parsed entry is too empty to be useful.
- Confidence is a number from 0 to 1.

Return this JSON shape:
{
  "review_items": [
    {
      "section": "experience",
      "status": "duplicate",
      "existing_index": 0,
      "parsed_index": 0,
      "confidence": 0.92,
      "reason": "Same company and similar role; parsed bullets differ.",
      "recommended_action": "merge"
    }
  ]
}

Existing profile sections:
%%EXISTING_PROFILE%%

Parsed profile sections:
%%PARSED_PROFILE%%
""".strip()


class ResumeParseError(RuntimeError):
    pass


def _clean_string(value) -> str:
    return str(value).strip() if value is not None else ""


def _normalize_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", _clean_string(value).lower()).strip()


def _file_extension(filename: str) -> str:
    lowered = (filename or "").lower()
    for extension in PDF_EXTENSIONS + DOCX_EXTENSIONS:
        if lowered.endswith(extension):
            return extension
    return ""


def is_supported_resume_file(uploaded_file) -> bool:
    content_type = (getattr(uploaded_file, "content_type", "") or "").lower()
    extension = _file_extension(getattr(uploaded_file, "name", ""))
    return (
        extension in PDF_EXTENSIONS
        or extension in DOCX_EXTENSIONS
        or content_type in PDF_CONTENT_TYPES
        or content_type in DOCX_CONTENT_TYPES
    )


def _is_docx_file(uploaded_file) -> bool:
    content_type = (getattr(uploaded_file, "content_type", "") or "").lower()
    extension = _file_extension(getattr(uploaded_file, "name", ""))
    return extension in DOCX_EXTENSIONS or content_type in DOCX_CONTENT_TYPES


def _extract_docx_text(uploaded_file) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeParseError("DOCX parsing requires python-docx to be installed.") from exc

    uploaded_file.seek(0)
    document = Document(BytesIO(uploaded_file.read()))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_resume_text(uploaded_file) -> str:
    if not is_supported_resume_file(uploaded_file):
        raise ResumeParseError("Resume must be a PDF or DOCX file.")

    if getattr(uploaded_file, "size", 0) > MAX_RESUME_BYTES:
        raise ResumeParseError("Resume file must be 10MB or smaller.")

    if _is_docx_file(uploaded_file):
        text = _extract_docx_text(uploaded_file)
    else:
        uploaded_file.seek(0)
        text = extract_pdf_text(uploaded_file)

    cleaned = re.sub(r"\n{3,}", "\n\n", (text or "").strip())
    if not cleaned:
        raise ResumeParseError("Could not extract readable text from the resume.")
    return cleaned


def _normalize_parsed_profile(payload: dict) -> dict:
    source = payload if isinstance(payload, dict) else {}
    normalized = normalize_profile_sections(source)
    for section in ("experience", "projects", "education", "certificates", "skills"):
        normalized[section] = [
            item
            for item in normalized.get(section, [])
            if any(_clean_string(value) for value in item.values() if not isinstance(value, (bool, list)))
            or any(_clean_string(entry) for value in item.values() if isinstance(value, list) for entry in value)
        ]
    normalized["full_name"] = _clean_string(source.get("full_name", ""))
    normalized["display_name"] = _clean_string(source.get("display_name", ""))
    return normalized


def _generate_profile_json(resume_text: str, provider: str, model: str) -> dict:
    prompt = PROFILE_PARSER_USER_PROMPT.replace("%%RESUME_TEXT%%", resume_text)
    if provider == "openai":
        response_text = _generate_with_openai(
            model=model,
            user_message=prompt,
            system_prompt=PROFILE_PARSER_SYSTEM_PROMPT,
        )
    elif provider == "anthropic":
        response_text = _generate_with_anthropic(
            model=model,
            user_message=prompt,
            system_prompt=PROFILE_PARSER_SYSTEM_PROMPT,
        )
    else:
        raise ResumeParseError(f"Unsupported provider: {provider}")

    parsed = _extract_json_object(response_text)
    if not parsed:
        try:
            parsed = json.loads(response_text)
        except json.JSONDecodeError as exc:
            raise ResumeParseError("Parser returned invalid JSON.") from exc

    if not isinstance(parsed, dict):
        raise ResumeParseError("Parser returned an unsupported JSON shape.")

    return _normalize_parsed_profile(parsed)


def _profile_sections_for_review(profile: dict) -> dict:
    normalized = _normalize_parsed_profile(profile)
    return {section: normalized.get(section, []) for section in PROFILE_REVIEW_SECTIONS}


def _generate_text(provider: str, model: str, prompt: str, system_prompt: str) -> str:
    if provider == "openai":
        return _generate_with_openai(
            model=model,
            user_message=prompt,
            system_prompt=system_prompt,
        )
    if provider == "anthropic":
        return _generate_with_anthropic(
            model=model,
            user_message=prompt,
            system_prompt=system_prompt,
        )
    raise ResumeParseError(f"Unsupported provider: {provider}")


def _extract_review_json(response_text: str) -> dict:
    parsed = _extract_json_object(response_text)
    if not parsed:
        try:
            parsed = json.loads(response_text)
        except json.JSONDecodeError as exc:
            raise ResumeParseError("Review returned invalid JSON.") from exc

    if not isinstance(parsed, dict):
        raise ResumeParseError("Review returned an unsupported JSON shape.")
    return parsed


def _coerce_index(value, section_length: int, field: str) -> int:
    if isinstance(value, bool) or not isinstance(value, int):
        raise ResumeParseError(f"Review item has invalid {field}.")
    if value < 0 or value >= section_length:
        raise ResumeParseError(f"Review item {field} is out of range.")
    return value


def _coerce_confidence(value) -> float:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        return 0.0
    return max(0.0, min(1.0, float(value)))


def _validate_review_items(payload: dict, existing_profile: dict, parsed_profile: dict) -> list[dict]:
    raw_items = payload.get("review_items")
    if not isinstance(raw_items, list):
        raise ResumeParseError("Review response is missing review_items.")

    existing_sections = _profile_sections_for_review(existing_profile)
    parsed_sections = _profile_sections_for_review(parsed_profile)
    expected = {
        (section, index)
        for section in PROFILE_REVIEW_SECTIONS
        for index, _item in enumerate(parsed_sections.get(section, []))
    }
    seen: set[tuple[str, int]] = set()
    review_items: list[dict] = []

    for raw_item in raw_items:
        if not isinstance(raw_item, dict):
            raise ResumeParseError("Review item has invalid shape.")

        section = _clean_string(raw_item.get("section", ""))
        if section not in PROFILE_REVIEW_SECTIONS:
            raise ResumeParseError("Review item has unsupported section.")

        parsed_index = _coerce_index(
            raw_item.get("parsed_index"),
            len(parsed_sections.get(section, [])),
            "parsed_index",
        )
        item_key = (section, parsed_index)
        if item_key in seen:
            raise ResumeParseError("Review response contains duplicate parsed entries.")
        seen.add(item_key)

        status = _clean_string(raw_item.get("status", ""))
        if status not in {"duplicate", "new"}:
            raise ResumeParseError("Review item has unsupported status.")

        existing_index = None
        if status == "duplicate":
            existing_index = _coerce_index(
                raw_item.get("existing_index"),
                len(existing_sections.get(section, [])),
                "existing_index",
            )
            allowed_actions = DUPLICATE_REVIEW_ACTIONS
            default_action = "merge"
        else:
            allowed_actions = NEW_REVIEW_ACTIONS
            default_action = "add"

        recommended_action = _clean_string(raw_item.get("recommended_action", ""))
        if recommended_action not in allowed_actions:
            recommended_action = default_action

        review_items.append({
            "section": section,
            "status": status,
            "existing_index": existing_index,
            "parsed_index": parsed_index,
            "confidence": _coerce_confidence(raw_item.get("confidence")),
            "reason": _clean_string(raw_item.get("reason", "")),
            "recommended_action": recommended_action,
        })

    if seen != expected:
        raise ResumeParseError("Review response did not cover every parsed entry.")

    return sorted(review_items, key=lambda item: (PROFILE_REVIEW_SECTIONS.index(item["section"]), item["parsed_index"]))


def _generate_review_items(existing_profile: dict, parsed_profile: dict, provider: str, model: str) -> list[dict]:
    prompt = (
        PROFILE_REVIEW_USER_PROMPT
        .replace("%%EXISTING_PROFILE%%", json.dumps(_profile_sections_for_review(existing_profile), ensure_ascii=False))
        .replace("%%PARSED_PROFILE%%", json.dumps(_profile_sections_for_review(parsed_profile), ensure_ascii=False))
    )
    response_text = _generate_text(provider, model, prompt, PROFILE_REVIEW_SYSTEM_PROMPT)
    return _validate_review_items(_extract_review_json(response_text), existing_profile, parsed_profile)


def _merge_fill_blanks(existing: dict, parsed: dict, fields: tuple[str, ...]) -> tuple[dict, list[dict]]:
    merged = dict(existing or {})
    suggestions: list[dict] = []

    for field in fields:
        current_value = _clean_string(merged.get(field, ""))
        parsed_value = _clean_string(parsed.get(field, ""))
        if not parsed_value:
            continue
        if not current_value:
            merged[field] = parsed_value
        elif current_value != parsed_value:
            suggestions.append({
                "section": "basic",
                "field": field,
                "current": current_value,
                "parsed": parsed_value,
            })

    return merged, suggestions


def _merge_string_list(existing: list, parsed: list, limit: int | None = None) -> list[str]:
    merged: list[str] = []
    seen: set[str] = set()
    for item in [*(existing or []), *(parsed or [])]:
        cleaned = _clean_string(item)
        key = _normalize_key(cleaned)
        if not cleaned or key in seen:
            continue
        seen.add(key)
        merged.append(cleaned)
        if limit is not None and len(merged) >= limit:
            break
    return merged


def _merge_items_by_key(
    existing_items: list[dict],
    parsed_items: list[dict],
    key_builder,
    section: str,
    bullet_field: str | None = None,
    list_merge_limits: dict[str, int | None] | None = None,
) -> tuple[list[dict], list[dict]]:
    merged = [dict(item) for item in existing_items if isinstance(item, dict)]
    matches: list[dict] = []
    index_by_key = {
        key_builder(item): index
        for index, item in enumerate(merged)
        if key_builder(item)
    }
    list_merge_limits = list_merge_limits or {}

    for parsed_item in parsed_items:
        if not isinstance(parsed_item, dict):
            continue

        parsed_key = key_builder(parsed_item)
        if parsed_key and parsed_key in index_by_key:
            item_index = index_by_key[parsed_key]
            current = dict(merged[item_index])
            for field, value in parsed_item.items():
                if field == bullet_field:
                    if isinstance(value, list) and value:
                        current[field] = value
                    continue
                if field in list_merge_limits:
                    current[field] = _merge_string_list(
                        current.get(field, []),
                        value if isinstance(value, list) else [],
                        limit=list_merge_limits[field],
                    )
                    continue
                if not _clean_string(current.get(field, "")) and _clean_string(value):
                    current[field] = value
            merged[item_index] = current
            matches.append({"section": section, "action": "updated", "key": parsed_key})
            continue

        merged.append(parsed_item)
        matches.append({"section": section, "action": "added", "key": parsed_key})

    return merged, matches


def _experience_key(item: dict) -> str:
    role = _normalize_key(item.get("role", ""))
    company = _normalize_key(item.get("company", ""))
    return f"{role}|{company}" if role and company else ""


def _project_key(item: dict) -> str:
    return _normalize_key(item.get("name", ""))


def _education_key(item: dict) -> str:
    school = _normalize_key(item.get("school", ""))
    degree = _normalize_key(item.get("degree", ""))
    major = _normalize_key(item.get("major", ""))
    return f"{school}|{degree}|{major}" if school else ""


def _certificate_key(item: dict) -> str:
    name = _normalize_key(item.get("name", ""))
    issuer = _normalize_key(item.get("issuer", ""))
    return f"{name}|{issuer}" if name else ""


def _skill_key(item: dict) -> str:
    return _normalize_key(item.get("name", ""))


def _fallback_review_items(existing_profile: dict, parsed_profile: dict) -> list[dict]:
    existing = _profile_sections_for_review(existing_profile)
    parsed = _profile_sections_for_review(parsed_profile)
    key_builders = {
        "experience": _experience_key,
        "projects": _project_key,
        "education": _education_key,
        "certificates": _certificate_key,
        "skills": _skill_key,
    }
    review_items: list[dict] = []

    for section in PROFILE_REVIEW_SECTIONS:
        key_builder = key_builders[section]
        index_by_key = {
            key_builder(item): index
            for index, item in enumerate(existing.get(section, []))
            if key_builder(item)
        }

        for parsed_index, parsed_item in enumerate(parsed.get(section, [])):
            parsed_key = key_builder(parsed_item)
            existing_index = index_by_key.get(parsed_key) if parsed_key else None
            if existing_index is None:
                review_items.append({
                    "section": section,
                    "status": "new",
                    "existing_index": None,
                    "parsed_index": parsed_index,
                    "confidence": 0.0,
                    "reason": "No deterministic match was found.",
                    "recommended_action": "add",
                })
                continue

            review_items.append({
                "section": section,
                "status": "duplicate",
                "existing_index": existing_index,
                "parsed_index": parsed_index,
                "confidence": 1.0,
                "reason": "Matched existing entry by normalized section key.",
                "recommended_action": "merge",
            })

    return review_items


def _merge_skills(existing_skills: list[dict], parsed_skills: list[dict]) -> tuple[list[dict], list[dict]]:
    merged = [dict(item) for item in existing_skills if isinstance(item, dict)]
    seen = {_normalize_key(item.get("name", "")) for item in merged if _normalize_key(item.get("name", ""))}
    matches: list[dict] = []

    for parsed_skill in parsed_skills:
        if not isinstance(parsed_skill, dict):
            continue
        key = _normalize_key(parsed_skill.get("name", ""))
        if not key or key in seen:
            continue
        seen.add(key)
        merged.append(parsed_skill)
        matches.append({"section": "skills", "action": "added", "key": key})

    return merged, matches


def merge_parsed_profile(existing_profile: dict, parsed_profile: dict) -> tuple[dict, list[dict], list[dict]]:
    existing = {
        **(existing_profile or {}),
        **normalize_profile_sections(existing_profile or {}),
    }
    parsed = _normalize_parsed_profile(parsed_profile)
    merged = dict(existing)
    suggestions: list[dict] = []
    matches: list[dict] = []

    for field in ("full_name", "display_name"):
        current_value = _clean_string(merged.get(field, ""))
        parsed_value = _clean_string(parsed.get(field, ""))
        if parsed_value and not current_value:
            merged[field] = parsed_value
        elif parsed_value and current_value and parsed_value != current_value:
            suggestions.append({
                "section": "profile",
                "field": field,
                "current": current_value,
                "parsed": parsed_value,
            })

    merged_basic, basic_suggestions = _merge_fill_blanks(
        existing.get("basic", {}),
        parsed.get("basic", {}),
        (
            "phone_country_code",
            "phone",
            "contact_email",
            "location",
            "headline",
            "github_url",
            "linkedin_url",
            "portfolio_url",
            "summary",
        ),
    )
    merged["basic"] = merged_basic
    suggestions.extend(basic_suggestions)

    merged["experience"], experience_matches = _merge_items_by_key(
        existing.get("experience", []),
        parsed.get("experience", []),
        _experience_key,
        "experience",
        bullet_field="description",
    )
    matches.extend(experience_matches)

    merged["projects"], project_matches = _merge_items_by_key(
        existing.get("projects", []),
        parsed.get("projects", []),
        _project_key,
        "projects",
        bullet_field="description",
        list_merge_limits={"stack": 5},
    )
    matches.extend(project_matches)

    merged["education"], education_matches = _merge_items_by_key(
        existing.get("education", []),
        parsed.get("education", []),
        _education_key,
        "education",
        list_merge_limits={"awards": None, "relevant_coursework": None},
    )
    matches.extend(education_matches)

    merged["certificates"], certificate_matches = _merge_items_by_key(
        existing.get("certificates", []),
        parsed.get("certificates", []),
        _certificate_key,
        "certificates",
    )
    matches.extend(certificate_matches)

    merged["skills"], skill_matches = _merge_skills(
        existing.get("skills", []),
        parsed.get("skills", []),
    )
    matches.extend(skill_matches)

    return _normalize_parsed_profile(merged) | {
        "uid": _clean_string(existing.get("uid", "")),
        "email": _clean_string(existing.get("email", "")),
        "photo_url": _clean_string(existing.get("photo_url", "")),
        "bio": _clean_string(existing.get("bio", "")),
        "created_at": _clean_string(existing.get("created_at", "")),
        "updated_at": _clean_string(existing.get("updated_at", "")),
    }, suggestions, matches


def parse_resume_into_profile(uploaded_file, existing_profile: dict, provider: str, model: str) -> dict:
    resume_text = extract_resume_text(uploaded_file)
    parsed_profile = _generate_profile_json(resume_text, provider, model)
    merged_profile, suggestions, matches = merge_parsed_profile(existing_profile, parsed_profile)
    parsed_sections = _profile_sections_for_review(parsed_profile)

    warnings = []
    if suggestions:
        warnings.append("Some parsed fields conflict with existing profile values and were left as suggestions.")

    if any(parsed_sections.get(section) for section in PROFILE_REVIEW_SECTIONS):
        try:
            review_items = _generate_review_items(existing_profile, parsed_profile, provider, model)
        except Exception:
            review_items = _fallback_review_items(existing_profile, parsed_profile)
            warnings.append("AI duplicate review was unavailable, so deterministic review matches were used.")
    else:
        review_items = []

    if not matches and not any(item["status"] == "duplicate" for item in review_items):
        warnings.append("No existing entries were matched; parsed entries were prepared as new profile data.")

    return {
        "parsed_profile": parsed_profile,
        "merged_profile": merged_profile,
        "suggestions": suggestions,
        "matches": matches,
        "review_items": review_items,
        "warnings": warnings,
    }
